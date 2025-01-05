import os
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import undetected_chromedriver as uc
import os
import docx
from docx import Document
from seleniumbase import Driver
from selenium.webdriver.common.action_chains import ActionChains

email=str(input('enter email'))
password12=str(input('enter password'))
project_name=str(input('enter name of the project'))


# Path to the folder containing videos
VIDEO_FOLDER_PATH = f"D:\\Automation PC\\{project_name}\\videos"
PHOTO_FOLDER_PATH = f"D:\\Automation PC\\{project_name}\\photos"
main_folder = f"D:\\Automation PC\\{project_name}"

# Function to get the list of video files in the folder
def get_video_files(folder_path):
    video_files = [f for f in os.listdir(folder_path) if f.endswith((".mp4", ".avi", ".mkv", ".mov"))]
    return video_files
def get_jpg_photos(folder_path):
    jpg_photos = glob.glob(os.path.join(folder_path, '*.jpg'))
    return jpg_photos

# Function to extract text from the first .docx file (for description)
def extract_text_from_first_docx(folder_path):
    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            docx_path = os.path.join(folder_path, file)
            doc = Document(docx_path)
            return "\n".join([para.text for para in doc.paragraphs])

# Function to extract text from the second .docx file (for tags)
def extract_text_from_second_docx(folder_path):
    docx_files = [file for file in os.listdir(folder_path) if file.endswith(".docx")]

    # Ensure there are at least two .docx files
    if len(docx_files) >= 2:
        second_docx_path = os.path.join(folder_path, docx_files[1])  # Get the second .docx file
        doc = Document(second_docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return "No second .docx file found for tags."
import glob

# Initialize ChromeDriver
driver = Driver(uc=True)

# Maximize the browser window
driver.maximize_window()

# Go to YouTube Studio
driver.get('https://studio.youtube.com')

# Wait for Google login to load and login
wait = WebDriverWait(driver, 6000)
email_input = wait.until(EC.presence_of_element_located((By.ID, "identifierId")))
email_input.send_keys(email)
email_input.send_keys(Keys.RETURN)

# Wait for the password page to load
time.sleep(8)  # adjust this based on your internet speed
password_input = wait.until(
    EC.presence_of_element_located((By.CSS_SELECTOR, "#password > div.aCsJod.oJeWuf > div > div.Xb9hP > input")))
password_input.send_keys(password12)
password_input.send_keys(Keys.RETURN)
time.sleep(120)
# Get the list of video files and description folders
video_files = [f for f in os.listdir(VIDEO_FOLDER_PATH) if f.endswith(".mp4")]
description_folders = [f for f in os.listdir(main_folder) if
                       os.path.isdir(os.path.join(main_folder, f))]
photo_files = get_jpg_photos(PHOTO_FOLDER_PATH)  # Get photos from the specific folder


# Ensure that the loop uses the same sequence for videos and descriptions
for video_file, description_folder, photo_file   in zip(video_files, description_folders, photo_files):
    video_path = os.path.join(VIDEO_FOLDER_PATH, video_file)
    folder_path = os.path.join(main_folder, description_folder)
    photo_path = os.path.join(PHOTO_FOLDER_PATH, photo_file)

    time.sleep(5)
    create_button = driver.find_element(By.ID, "create-icon")
    create_button.click()
    time.sleep(2)
    # Click on the 'Upload videos' option
    upload_videos = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#text-item-0")))
    upload_videos.click()
    time.sleep(3)
    # Upload the video file
    file_input = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@type='file']")))
    file_input.send_keys(video_path)
    time.sleep(10)
    textbox = wait.until(EC.presence_of_element_located((By.ID, "textbox")))
    # Clear the existing text
    textbox.clear()
    time.sleep(3)
    video_name = video_file[:-4]

    # Input your new text
    textbox.send_keys(video_name)
    time.sleep(5)
    textbox.clear()
    textbox.send_keys(video_name)
    # Extract description from the corresponding folder
    description_text = extract_text_from_first_docx(folder_path)
    # Use the extracted text in the description field if available
    if description_text:
        description_field = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#description-textarea")))
        description_field.send_keys(description_text)
    else:
        print(f"No .docx file found in {folder_path}")
    time.sleep(4)
    file_input = driver.find_element(By.CSS_SELECTOR, "input[type='file']")
    file_input.send_keys(photo_path)
    radio_button = wait.until(
        EC.presence_of_element_located(
            (By.CSS_SELECTOR, 'tp-yt-paper-radio-button[name="VIDEO_MADE_FOR_KIDS_NOT_MFK"]')))
    # If found, click the radio button
    radio_button.click()
    time.sleep(3)

    show_more_button = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#toggle-button > ytcp-button-shape > button')))
    show_more_button.click()

    altered_button = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR,
                                                                '#details > div > ytcp-video-metadata-editor-advanced > div:nth-child(4) > ytkp-altered-content-select > div.altered-content-rating-container.style-scope.ytkp-altered-content-select > tp-yt-paper-radio-group > tp-yt-paper-radio-button:nth-child(1)')))
    altered_button.click()
    time.sleep(3)
    tags_text = extract_text_from_second_docx(folder_path)
    tags = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '#text-input')))
    tags.send_keys(tags_text)
    time.sleep(3)
    next_button = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#next-button > ytcp-button-shape > button')))
    next_button.click()
    time.sleep(2)




    #monetization page
    arrow = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '.ytcp-video-monetization .edit-button.show'))
    )
    arrow.click()
    time.sleep(5)
    on_button = wait.until(
        EC.presence_of_element_located((By.XPATH, '//*[@id="radio-on"]'))
    )
    on_button.click()
    time.sleep(5)
    done_button = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#save-button > ytcp-button-shape:nth-child(1)'))
    )
    done_button.click()
    next_button2 = wait.until(
        EC.presence_of_element_located(
            (By.CSS_SELECTOR, '#next-button > ytcp-button-shape:nth-child(1) > button:nth-child(1)'))
    )
    next_button2.click()


    #ads page
    # None of the Above Option
    none_of_the_above_option = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '.all-none-checkbox'))
    )
    none_of_the_above_option.click()
    time.sleep(5)
    # Submit Rating Button
    submit_rating_button = wait.until(
        EC.presence_of_element_located(
            (By.CSS_SELECTOR, '#submit-questionnaire-button > ytcp-button-shape:nth-child(1) > button:nth-child(1)'))
    )
    submit_rating_button.click()
    time.sleep(15)

    # Last Next Button
    next_button3 = wait.until(
        EC.presence_of_element_located(
            (By.CSS_SELECTOR, '#next-button > ytcp-button-shape:nth-child(1) > button:nth-child(1)'))
    )
    next_button3.click()

    #video elements page
    time.sleep(8)
    import_screens_button = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#import-from-video-button > ytcp-button-shape > button')))
    import_screens_button.click()
    time.sleep(5)
    sepc_import = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR,'#import-endscreen-from-video-button > ytcp-button-shape > button')))
    sepc_import.click()
    time.sleep(8)
    last_card = None
    while True:
        try:
            # Scroll down by simulating PAGE_DOWN key presses
            driver.find_element(By.TAG_NAME, "body").send_keys(Keys.PAGE_DOWN)
            time.sleep(1)  # Adjust sleep time if needed

            # Find all video cards by their class name
            video_cards = driver.find_elements(By.CSS_SELECTOR, "ytcp-entity-card")

            # If all cards are loaded, exit the loop
            if len(video_cards) > 0:
                last_card = video_cards[-1]  # Get the last video card
                break
        except Exception as e:
            print(f"Error while scrolling: {e}")
            break

    # Scroll the last card into view and click it
    driver.execute_script("arguments[0].scrollIntoView();", last_card)
    time.sleep(1)  # Wait for it to scroll into view

    # Click the last video card
    try:
        last_card.click()
        print("Successfully clicked the last video card!")
    except Exception as e:
        print(f"Error clicking the last card: {e}")
    time.sleep(7)
    save_screen= wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#save-button > ytcp-button-shape > button')))
    save_screen.click()
    time.sleep(4)
    navigate_to_visbility = wait.until(
        EC.presence_of_element_located((By.XPATH, '//*[@id="step-badge-3"]')))
    navigate_to_visbility.click()
    time.sleep(2)

    #Visibilty page
    public_radio_button = wait.until(EC.element_to_be_clickable((By.NAME, "PUBLIC")))
    # Click the radio button
    public_radio_button.click()
    time.sleep(2)

    next_button6 = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#done-button > ytcp-button-shape > button')))
    next_button6.click()
    time.sleep(7)
    buttons = driver.find_elements(By.CSS_SELECTOR, '#secondary-action-button > ytcp-button-shape > button')
    # If the button exists (list is not empty), click it
    if buttons:
        buttons[0].click()  # Click the first (and only) button
        print("Button clicked!")
    else:
        print("Button not found, skipping.")
    time.sleep(4)
    close_button = wait.until(
        EC.presence_of_element_located((By.CSS_SELECTOR, '#close-button > ytcp-button-shape > button')))
    close_button.click()

    time.sleep(5)











